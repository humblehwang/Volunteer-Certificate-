# -*- coding: utf-8 -*-
"""
Spyder Editor

This is a temporary script file.
"""



import pandas as pd
import xlrd
import docx 
from docx.enum.style import WD_STYLE_TYPE

doc_new = docx.Document()
doc = docx.Document("D:/NCTU/2018Summer/fen/example.docx")


parag_num = 0
str_word = ""
for para in doc.paragraphs :
    str_word += para.text



data_xls = pd.read_excel("D:/NCTU/2018Summer/fen/time.xlsx", index_col=0)
data_xls.to_csv('time.csv', encoding='utf-8')
for index in range(len(data_xls)):
    
    str1 = "志工證明"
    str2 = "東華大學"
    str3 = "於中華民國一百零七年三月至五月間，參與志工服務訓練，協助"
    str4 = "，共"
    str5 = "小時，特頒此狀，以表感謝。"
    str6 = "中華民國107年1月7日"
    output = ""
   
    output =  str2 + data_xls['department'][index] + data_xls['name'][index] + str3 + data_xls['item'][index] + str4 + str(data_xls['time'][index]) + str5        
    file = 'D:/NCTU/2018Summer/fen/' + data_xls['name'][index] + ".docx"
    
    
    
    doc_output = docx.Document()
    style =  doc_output.styles['Normal']
    font = style.font


    from docx.shared import Pt
    font.name = "華康正顏楷體W9"
    font.size = Pt(30)
    
    
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    paragraph = doc_output.add_paragraph(str1)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc_output.add_paragraph(output)
    doc_output.add_paragraph(str6)
    
    
  


    doc_output.save(file)

